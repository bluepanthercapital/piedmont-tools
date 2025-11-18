import streamlit as st
import pandas as pd
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


LOGO_PATH = "piedmont_logo.png"


# Helper to clean NaN -> ""
def fmt(val):
    if pd.isna(val):
        return ""
    return str(val)


# Helper: shade + bold header row
def style_header_row(row):
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "DDDDDD")  # light grey
        tc_pr.append(shd)

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)


# Columns we absolutely need (Kiosk handled separately)
REQUIRED_COLUMNS = [
    "Property Name",
    "Store ID (NSA Unique ID)",
    "Address",
    "City",
    "St",
    "Zip",
    "Size",
    "Gen",
    "Indoor/ Outdoor",
    "Config",
    "Locker Name",
    "Contact Name",
    "Contact Phone #",
    "PO for Invoice",
]


# ---------------- Streamlit UI ----------------

st.set_page_config(page_title="Amazon Locker Sheet Generator", layout="centered")

# Top bar: logo + title
if Path(LOGO_PATH).exists():
    col_logo, col_title = st.columns([1, 3])
    with col_logo:
        st.image(LOGO_PATH, use_column_width="always")
    with col_title:
        st.title("Amazon Locker Sheet Generator")
else:
    st.title("Amazon Locker Sheet Generator")

st.write(
    "Upload the Amazon Locker Excel file, select a Locker Name, "
    "and generate a Word document with the property and locker details."
)

uploaded_file = st.file_uploader("Upload Excel file (.xlsx)", type=["xlsx"])

if uploaded_file is not None:
    # Read first sheet by default
    try:
        xls = pd.ExcelFile(uploaded_file)
        sheet_name = xls.sheet_names[0]
        df = pd.read_excel(xls, sheet_name=sheet_name)
    except Exception as e:
        st.error(f"Error reading Excel file: {e}")
        st.stop()

    # Check that required columns exist (excluding Kiosk for now)
    missing = [c for c in REQUIRED_COLUMNS if c not in df.columns]
    if missing:
        st.error(
            "The uploaded file is missing these columns: "
            + ", ".join(missing)
        )
        st.stop()

    # Handle Kiosk column which may be named "Kiosk " (with space) or "Kiosk"
    if "Kiosk " in df.columns:
        kiosk_col = "Kiosk "
    elif "Kiosk" in df.columns:
        kiosk_col = "Kiosk"
    else:
        st.error(
            "The uploaded file is missing the 'Kiosk' column "
            "(expected header 'Kiosk' or 'Kiosk ')."
        )
        st.stop()

    # Drop rows where Locker Name is missing
    locker_series = df["Locker Name"].dropna().astype(str)
    locker_options = sorted(locker_series.unique())

    if not locker_options:
        st.error("No Locker Name values found in the file.")
        st.stop()

    selected_locker = st.selectbox(
        "Select Locker Name",
        options=locker_options,
        index=0,
    )

    # Filter rows for selected locker
    matching_rows = df[df["Locker Name"].astype(str) == selected_locker]

    st.subheader("Row preview")
    preview_cols = [
        "Locker Name",
        kiosk_col,
        "Property Name",
        "Store ID (NSA Unique ID)",
        "Address",
        "City",
        "St",
        "Zip",
        "Size",
        "Gen",
        "Indoor/ Outdoor",
        "Config",
        "Contact Name",
        "Contact Phone #",
        "PO for Invoice",
    ]
    st.dataframe(matching_rows[preview_cols])

    st.write(
        "If there are multiple rows with the same Locker Name, "
        "the first one will be used for the Word document for now."
    )

    if st.button("Generate Word document"):
        if matching_rows.empty:
            st.error("No row found for that Locker Name.")
            st.stop()

        row = matching_rows.iloc[0]

        # ---------------- Word document ----------------
        doc = Document()

        # Page margins
        section = doc.sections[0]
        section.top_margin = Inches(1.1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

        # Header with logo (top-right)
        if Path(LOGO_PATH).exists():
            header = section.header
            header_para = (
                header.paragraphs[0]
                if header.paragraphs
                else header.add_paragraph()
            )
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            header_run = header_para.add_run()
            try:
                header_run.add_picture(LOGO_PATH, width=Inches(2.5))
            except Exception:
                # If anything goes wrong with the image, just skip it
                pass

        # Footer with small Piedmont line
        footer = section.footer
        footer_para = (
            footer.paragraphs[0]
            if footer.paragraphs
            else footer.add_paragraph()
        )
        footer_para.text = (
            "Piedmont Moving Systems • Amazon Locker Operations"
        )
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_para.runs:
            run.font.size = Pt(8)
            run.font.italic = True

        # Title: Locker name + kiosk
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(
            f"{fmt(row['Locker Name'])}  (Kiosk {fmt(row[kiosk_col])})"
        )
        title_run.bold = True
        title_run.font.size = Pt(24)

        doc.add_paragraph()  # spacer

        # ---- Property Details header ----
        prop_header = doc.add_paragraph()
        prop_header_run = prop_header.add_run("Property Details")
        prop_header_run.bold = True
        prop_header_run.font.size = Pt(14)

        doc.add_paragraph()  # small spacer

        # ---- Table 1: Property block ----
        table1 = doc.add_table(rows=2, cols=6)
        table1.style = "Table Grid"
        table1.autofit = False

        col_widths1 = [
            Inches(2.0),
            Inches(1.8),
            Inches(2.2),
            Inches(1.3),
            Inches(0.7),
            Inches(0.9),
        ]
        for col, width in zip(table1.columns, col_widths1):
            col.width = width

        hdr_cells = table1.rows[0].cells
        hdr_cells[0].text = "Property Name"
        hdr_cells[1].text = "Store ID (NSA Unique ID)"
        hdr_cells[2].text = "Address"
        hdr_cells[3].text = "City"
        hdr_cells[4].text = "St"
        hdr_cells[5].text = "Zip"
        style_header_row(table1.rows[0])

        val_cells = table1.rows[1].cells
        val_cells[0].text = fmt(row["Property Name"])
        val_cells[1].text = fmt(row["Store ID (NSA Unique ID)"])
        val_cells[2].text = fmt(row["Address"])
        val_cells[3].text = fmt(row["City"])
        val_cells[4].text = fmt(row["St"])
        val_cells[5].text = fmt(row["Zip"])

        # Adjust body font
        for row_cells in table1.rows:
            for cell in row_cells.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if not run.bold:
                            run.font.size = Pt(11)
                # vertical center
                cell.vertical_alignment = 1  # 0=top,1=center,2=bottom

        doc.add_paragraph()  # spacer

        # ---- Locker Details header ----
        lock_header = doc.add_paragraph()
        lock_header_run = lock_header.add_run("Locker Details")
        lock_header_run.bold = True
        lock_header_run.font.size = Pt(14)

        doc.add_paragraph()  # small spacer

        # ---- Table 2: Locker and contact block ----
        table2 = doc.add_table(rows=2, cols=8)
        table2.style = "Table Grid"
        table2.autofit = False

        col_widths2 = [
            Inches(0.7),
            Inches(0.8),
            Inches(1.3),
            Inches(1.0),
            Inches(1.8),
            Inches(1.8),
            Inches(1.6),
            Inches(1.8),
        ]
        for col, width in zip(table2.columns, col_widths2):
            col.width = width

        hdr2 = table2.rows[0].cells
        hdr2[0].text = "Size"
        hdr2[1].text = "Gen"
        hdr2[2].text = "Indoor/ Outdoor"
        hdr2[3].text = "Config"
        hdr2[4].text = "Locker Name"
        hdr2[5].text = "Contact Name"
        hdr2[6].text = "Contact Phone #"
        hdr2[7].text = "PO for Invoice"
        style_header_row(table2.rows[0])

        val2 = table2.rows[1].cells
        val2[0].text = fmt(row["Size"])
        val2[1].text = fmt(row["Gen"])
        val2[2].text = fmt(row["Indoor/ Outdoor"])
        val2[3].text = fmt(row["Config"])
        val2[4].text = fmt(row["Locker Name"])
        val2[5].text = fmt(row["Contact Name"])
        val2[6].text = fmt(row["Contact Phone #"])
        val2[7].text = fmt(row["PO for Invoice"])

        for row_cells in table2.rows:
            for cell in row_cells.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if not run.bold:
                            run.font.size = Pt(11)
                cell.vertical_alignment = 1

        # Save to memory buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        file_name = f"{fmt(row['Locker Name'])}_Kiosk_{fmt(row[kiosk_col])}.docx"

        st.download_button(
            label="Download Word document",
            data=buffer,
            file_name=file_name,
            mime=(
                "application/"
                "vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
        )
